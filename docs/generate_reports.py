#!/usr/bin/env python3
"""Generate enterprise-grade PostureOne reports (PDF + DOCX)."""

from __future__ import annotations

import os
from datetime import datetime

# ─── Report 1: Product Overview PDF ─────────────────────────────────

def generate_product_pdf(output_path: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm, cm
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, HRFlowable,
    )

    BRAND = HexColor("#6366f1")
    DARK = HexColor("#1e1b4b")
    GRAY = HexColor("#475569")
    LIGHT_BG = HexColor("#f8fafc")
    WHITE = HexColor("#ffffff")
    RED = HexColor("#dc2626")
    GREEN = HexColor("#16a34a")
    ORANGE = HexColor("#ea580c")

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        leftMargin=25*mm, rightMargin=25*mm,
        topMargin=30*mm, bottomMargin=25*mm,
        title="PostureOne - Product Overview",
        author="Andrea Ceresoni",
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        "CoverTitle", parent=styles["Title"],
        fontSize=28, textColor=DARK, spaceAfter=6,
        alignment=TA_CENTER, fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        "CoverSub", parent=styles["Normal"],
        fontSize=14, textColor=BRAND, spaceAfter=30,
        alignment=TA_CENTER, fontName="Helvetica",
    ))
    styles.add(ParagraphStyle(
        "SectionH", parent=styles["Heading1"],
        fontSize=18, textColor=DARK, spaceBefore=20, spaceAfter=10,
        fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        "SubH", parent=styles["Heading2"],
        fontSize=13, textColor=BRAND, spaceBefore=14, spaceAfter=6,
        fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontSize=10, textColor=GRAY, leading=15,
        alignment=TA_JUSTIFY, fontName="Helvetica",
    ))
    styles.add(ParagraphStyle(
        "BodyBold", parent=styles["Normal"],
        fontSize=10, textColor=DARK, leading=15,
        fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        "SmallGray", parent=styles["Normal"],
        fontSize=8, textColor=GRAY, alignment=TA_CENTER,
    ))
    styles.add(ParagraphStyle(
        "CenterBody", parent=styles["Normal"],
        fontSize=10, textColor=GRAY, alignment=TA_CENTER,
    ))

    story = []

    def hr():
        story.append(Spacer(1, 4))
        story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#e2e8f0")))
        story.append(Spacer(1, 8))

    def section(title):
        story.append(Paragraph(title, styles["SectionH"]))
        hr()

    def sub(title):
        story.append(Paragraph(title, styles["SubH"]))

    def body(text):
        story.append(Paragraph(text, styles["Body"]))
        story.append(Spacer(1, 6))

    def bold(text):
        story.append(Paragraph(text, styles["BodyBold"]))
        story.append(Spacer(1, 4))

    def make_table(data, col_widths=None):
        t = Table(data, colWidths=col_widths, hAlign="LEFT")
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), BRAND),
            ("TEXTCOLOR", (0, 0), (-1, 0), WHITE),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 9),
            ("FONTSIZE", (0, 1), (-1, -1), 9),
            ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
            ("TEXTCOLOR", (0, 1), (-1, -1), DARK),
            ("BACKGROUND", (0, 1), (-1, -1), LIGHT_BG),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [LIGHT_BG, WHITE]),
            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#e2e8f0")),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ]))
        story.append(t)
        story.append(Spacer(1, 12))

    # ── COVER ──
    story.append(Spacer(1, 60))
    story.append(Paragraph("PostureOne", styles["CoverTitle"]))
    story.append(Paragraph("Cloud Security Posture Management Platform", styles["CoverSub"]))
    story.append(Spacer(1, 20))
    story.append(HRFlowable(width="40%", thickness=2, color=BRAND))
    story.append(Spacer(1, 20))
    story.append(Paragraph("Product Overview &amp; Technical Capabilities", styles["CenterBody"]))
    story.append(Spacer(1, 40))
    story.append(Paragraph("Andrea Ceresoni", styles["CenterBody"]))
    story.append(Paragraph(f"April 2026  |  Confidential", styles["SmallGray"]))
    story.append(PageBreak())

    # ── EXECUTIVE SUMMARY ──
    section("Executive Summary")
    body(
        "PostureOne is a multi-tenant Cloud Security Posture Management (CSPM) platform "
        "that continuously scans cloud infrastructure, evaluates it against 179 security checks "
        "mapped to CIS Benchmarks v3.0, and delivers a prioritised action plan that tells "
        "security teams exactly what to fix first and why."
    )
    body(
        "Unlike competitors that use proprietary black-box ML scoring (Wiz, Prisma Cloud) or "
        "produce flat, unprioritised finding lists (Prowler, ScoutSuite), PostureOne uses a "
        "transparent, YAML-driven priority formula that buyers can inspect, customise, and audit. "
        "This positions PostureOne uniquely at the intersection of open-source transparency and "
        "enterprise-grade triage."
    )

    # ── THE PROBLEM ──
    section("The Problem")
    body(
        "Every CSPM tool generates findings. A typical Azure subscription with 50-200 resources "
        "produces 60-200 failing controls on day one. Without triage, the user sees a flat list "
        "of red items with no indication of which ones matter, which ones are quick to fix, or "
        "which ones are internet-exposed. The result is predictable: the user opens the findings "
        "page, feels overwhelmed, closes it, and nothing gets fixed."
    )
    body(
        "This is the core retention problem in CSPM: the tool produces signal, but the user "
        "cannot act on it because the signal-to-noise ratio is too low. Alerting without triage "
        "creates noise. Dashboards without prioritisation create confusion."
    )

    # ── THE SOLUTION ──
    section("The Solution: Transparent Priority Triage")
    body(
        "PostureOne automatically assigns every failing finding a priority bucket (P0 / P1 / P2 / P3) "
        "computed from three transparent axes:"
    )
    make_table([
        ["Axis", "Values", "Source"],
        ["Severity", "high / medium / low", "CIS benchmark definition"],
        ["Effort", "quick / moderate / refactor", "YAML or inferred from control name"],
        ["Exposure", "internet / internal / none", "YAML or inferred from resource type"],
    ], col_widths=[80, 130, 230])

    body("The formula is a 3x3 matrix with an internet-exposure bump:")
    make_table([
        ["", "Quick", "Moderate", "Refactor"],
        ["High severity", "P0", "P1", "P2"],
        ["Medium severity", "P1", "P2", "P3"],
        ["Low severity", "P2", "P3", "P3"],
    ], col_widths=[110, 100, 100, 100])
    body(
        "Internet-exposed findings are bumped up one tier (capped at P0). The formula, the "
        "defaults, and the per-control overrides all live in control_mappings.yaml and are "
        "100% inspectable and customisable per tenant."
    )

    # ── KEY DIFFERENTIATORS ──
    section("Key Differentiators")
    make_table([
        ["Capability", "PostureOne", "Wiz / Prisma", "Prowler / ScoutSuite"],
        ["Priority triage", "Transparent YAML formula", "Black-box ML", "None (flat list)"],
        ["Remediation snippets", "40 auto-filled (Terraform/Bicep/CLI)", "Generic", "None"],
        ["Secure Score projection", "Current / after P0 / after P0+P1", "Single score", "None"],
        ["Multi-tenant", "First-class (MSSP-ready)", "Enterprise add-on", "No"],
        ["Customisable per tenant", "Full YAML override", "No", "N/A"],
        ["Smart alerting", "Only NEW P0/P1 (delta-based)", "All findings", "None"],
        ["Scan drill-down", "Delta vs previous scan", "Yes", "No"],
        ["Pricing", "Transparent tiers", "Enterprise quotes", "Free (no support)"],
    ], col_widths=[110, 110, 110, 110])

    # ── PLATFORM CAPABILITIES ──
    section("Platform Capabilities")

    sub("Security Checks")
    body("179 built-in security checks across 46 resource types, mapped to CIS Benchmarks v3.0:")
    make_table([
        ["Provider", "Evaluators", "YAML Controls", "Resource Types"],
        ["Azure", "159", "163", "34"],
        ["AWS", "20", "20", "12"],
        ["Total", "179", "183", "46"],
    ], col_widths=[100, 100, 120, 120])

    sub("Remediation Auto-Fill")
    body(
        "40 controls ship with IaC remediation snippets in three formats (Terraform, Bicep, Azure CLI). "
        "Snippets are rendered with the actual asset's subscription ID, resource group, and resource name "
        "extracted from the ARM ID -- not generic placeholders. The user clicks Copy and pastes a "
        "ready-to-run command."
    )

    sub("Smart Alerting")
    body(
        "The finding.new_p0 event fires only when NEW P0 or P1 findings appear that were not present "
        "in the previous scan. This prevents alert fatigue. Delivered via Slack webhooks with priority "
        "emoji, Secure Score, delta counts, and action buttons linking directly to the findings page."
    )

    sub("Additional Capabilities")
    make_table([
        ["Feature", "Description"],
        ["Scan drill-down", "Priority breakdown + delta (new/fixed/unchanged) vs previous scan"],
        ["Secure Score projection", "Current / after P0 / after P0+P1 visualised as a segmented bar"],
        ["Grouped remediation", "19 remediation groups aggregate controls with the same fix"],
        ["Bulk operations", "Waive, assign, comment on findings at scale"],
        ["PDF evidence packs", "One-click export for audits and compliance reports"],
        ["SSO/OIDC + MFA", "Azure AD, Okta, Google Workspace + TOTP with backup codes"],
        ["Custom RBAC", "Granular roles beyond admin/viewer"],
        ["Scheduled scans", "Cron-based automated scanning via Celery Beat"],
        ["Jira integration", "Push findings as tickets"],
        ["API keys", "Programmatic access for CI/CD pipelines"],
    ], col_widths=[130, 310])

    story.append(PageBreak())

    # ── ARCHITECTURE ──
    section("Architecture Overview")
    make_table([
        ["Layer", "Technology"],
        ["Frontend", "Next.js 16 (App Router), TypeScript, Tailwind CSS, Recharts"],
        ["Backend", "Python 3.12, FastAPI, SQLAlchemy 2.x (async), Pydantic v2"],
        ["Database", "PostgreSQL 16"],
        ["Cache / Queue", "Redis 7, Celery"],
        ["Authentication", "JWT (httpOnly cookies), bcrypt, TOTP/MFA, SSO/OIDC"],
        ["Infrastructure", "Docker Compose, Caddy (TLS), GitHub Actions CI/CD"],
    ], col_widths=[100, 340])

    # ── LIVE RESULTS ──
    section("Live Validation: IFO Production Tenant")
    body(
        "The following results were captured on a live Azure tenant (6 assets, 86 evaluations) "
        "running the full 163-control scan against a real production subscription:"
    )
    make_table([
        ["Priority Bucket", "Count", "Action Required"],
        ["P0 -- Fix now", "25", "High severity + quick fix + internet-exposed"],
        ["P1 -- Fix this week", "29", "Structural hardening"],
        ["P2 -- Fix this sprint", "8", "Governance and resilience"],
        ["P3 -- Best practice", "3", "Polish when time permits"],
        ["Total failing", "65", ""],
        ["Passing", "21", ""],
    ], col_widths=[140, 60, 240])

    body(
        "Secure Score: 24.4% (current) -> 53.5% (after fixing P0) -> 87.2% (after fixing P0 + P1). "
        "This demonstrates that one week of focused P0 remediation more than doubles the security posture."
    )

    # ── SECURITY POSTURE ──
    section("Platform Security Posture")
    body(
        "PostureOne underwent a comprehensive 4-reviewer parallel audit in April 2026 covering "
        "security, code quality, architecture, and product UX."
    )
    make_table([
        ["Area", "Grade", "Key Findings"],
        ["Security", "A-", "0 Critical, 0 High vulnerabilities. SSRF protection, JWT production-grade"],
        ["Code Quality", "B", "1297 backend tests, evaluator batched, cursor pagination"],
        ["Architecture", "B+", "Tenant isolation verified (322 JOIN checks), NOT NULL enforcement"],
        ["Product", "B", "Triage layer unique in market, 40 auto-fill snippets, scan drill-down"],
        ["Dependencies", "A-", "1 moderate Dependabot alert (from 11)"],
    ], col_widths=[80, 40, 320])

    # ── PRICING ──
    section("Pricing Tiers")
    make_table([
        ["Tier", "Capabilities", "Target Price Range"],
        ["Scanner", "Asset inventory, security checks, Secure Score, PDF export", "5,000 - 15,000 EUR/year"],
        ["Advisor", "Scanner + Priority triage, auto-fill remediation, smart alerting, scan delta", "40,000 - 100,000 EUR/year"],
    ], col_widths=[70, 260, 110])
    body(
        "The Advisor tier is unlocked by the Priority/Triage layer -- the same underlying checks, "
        "reframed as a ranked action plan with projected impact. This is a 3-5x price multiplier "
        "without adding any new scanning capability."
    )

    # ── FOOTER ──
    story.append(Spacer(1, 30))
    story.append(HRFlowable(width="100%", thickness=1, color=BRAND))
    story.append(Spacer(1, 8))
    story.append(Paragraph(
        "Copyright 2026 Andrea Ceresoni. All rights reserved. Confidential.",
        styles["SmallGray"],
    ))

    doc.build(story)
    print(f"  PDF generated: {output_path}")


# ─── Report 2: Technical Assessment DOCX ────────────────────────────

def generate_technical_docx(output_path: str) -> None:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    BRAND = RGBColor(0x63, 0x66, 0xf1)
    DARK = RGBColor(0x1e, 0x1b, 0x4b)
    GRAY = RGBColor(0x47, 0x55, 0x69)

    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = DARK if level == 1 else BRAND
        return h

    def add_body(text):
        p = doc.add_paragraph(text)
        p.style.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_table(headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Header row
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        # Data rows
        for r, row in enumerate(rows):
            for c, val in enumerate(row):
                cell = table.rows[r + 1].cells[c]
                cell.text = str(val)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

        doc.add_paragraph("")
        return table

    # ── COVER ──
    doc.add_paragraph("")
    doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PostureOne")
    run.font.size = Pt(32)
    run.font.color.rgb = DARK
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Technical Assessment Report")
    run.font.size = Pt(16)
    run.font.color.rgb = BRAND

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("Architecture, Security & Code Quality Review - April 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY

    doc.add_paragraph("")
    doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Author: ").font.color.rgb = GRAY
    run = meta.add_run("Andrea Ceresoni")
    run.bold = True
    run.font.color.rgb = DARK

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta2.add_run("Classification: Confidential")
    run.font.color.rgb = GRAY
    run.font.size = Pt(9)

    doc.add_page_break()

    # ── TABLE OF CONTENTS (manual) ──
    add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Platform Architecture",
        "3. Security Assessment",
        "4. Code Quality Assessment",
        "5. Scalability Assessment",
        "6. Test Coverage",
        "7. Dependency Health",
        "8. Production Deployment",
        "9. Recommendations",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()

    # ── 1. EXECUTIVE SUMMARY ──
    add_heading("1. Executive Summary", level=1)
    add_body(
        "This report documents the results of a comprehensive technical assessment of the "
        "PostureOne CSPM platform, conducted in April 2026 by four independent reviewers "
        "covering security, code quality, architecture, and product readiness."
    )
    add_body("Overall assessment grades:")
    add_table(
        ["Area", "Grade", "Summary"],
        [
            ["Security", "A-", "0 Critical, 0 High. All OWASP Top 10 addressed. SSRF double-layer, JWT production-grade."],
            ["Code Quality", "B", "1297 backend tests. Evaluator batched, cursor pagination, seed upsert atomic."],
            ["Architecture", "B+", "Tenant isolation verified across 322 query join points. NOT NULL enforced."],
            ["Product / UX", "B", "Unique triage layer, 40 auto-fill remediation snippets, scan drill-down with delta."],
            ["Dependencies", "A-", "Dependabot alerts reduced from 11 to 1 (moderate). All critical/high resolved."],
        ],
    )
    add_body(
        "The platform was elevated from C+/B- to B+/A- over a 3-day intensive sprint that "
        "addressed all Critical and High findings, implemented cursor pagination, batched the "
        "evaluator for large subscriptions, and added smart alerting on new P0 findings."
    )

    # ── 2. ARCHITECTURE ──
    add_heading("2. Platform Architecture", level=1)
    add_body(
        "PostureOne follows a standard three-tier architecture with asynchronous task processing "
        "for long-running cloud scans."
    )
    add_table(
        ["Component", "Technology", "Purpose"],
        [
            ["API Server", "FastAPI (Python 3.12, async)", "RESTful API with JWT auth, RBAC, rate limiting"],
            ["Frontend", "Next.js 16 (App Router, TypeScript)", "Server-rendered React UI with Tailwind CSS"],
            ["Database", "PostgreSQL 16", "10 core tables, Alembic migrations, async driver (asyncpg)"],
            ["Cache / Broker", "Redis 7", "Celery task broker, API response cache, rate limiter backend"],
            ["Worker", "Celery (prefork, concurrency=2)", "Scan pipeline: collector -> normalizer -> evaluator"],
            ["Reverse Proxy", "Caddy", "Automatic TLS via Let's Encrypt, HTTP/2"],
            ["CI/CD", "GitHub Actions", "Test + Trivy scan + Docker build + SSH deploy"],
        ],
    )

    add_heading("Data Pipeline", level=2)
    add_body(
        "The scan pipeline executes as a single Celery task with three sequential stages: "
        "(1) Collector queries Azure Resource Graph and Defender APIs to build the asset inventory; "
        "(2) Normalizer maps Defender recommendations to internal controls via dedup keys; "
        "(3) Evaluator runs 179 registered check functions against each asset and creates/updates findings. "
        "Assets are processed in batches of 500 to prevent OOM on large subscriptions."
    )

    add_heading("Multi-Tenancy Model", level=2)
    add_body(
        "Tenant isolation is enforced at the query level via JOIN on cloud_accounts.tenant_id. "
        "This pattern was verified across 322 occurrences in all 26 API endpoint files. "
        "The tenant_id column on findings and assets is now NOT NULL with a foreign key to the "
        "tenants table, backed by an Alembic migration that backfilled all existing rows."
    )

    # ── 3. SECURITY ──
    add_heading("3. Security Assessment", level=1)
    add_body(
        "The security assessment identified 3 Critical, 8 High, and 8 Medium findings. "
        "All Critical and High findings have been remediated."
    )

    add_heading("Resolved Findings", level=2)
    add_table(
        ["ID", "Severity", "Finding", "Resolution"],
        [
            ["C-1", "Critical", "Password reset token logged in plaintext", "Token removed from log output"],
            ["C-2", "Critical", "Fernet encryption key hardcoded in CI workflow", "Fallback removed, requires GitHub Secret"],
            ["C-3", "Critical", "Open redirect on login page via redirect param", "Validated relative path only"],
            ["H-1", "High", "SSRF via OIDC discovery (plain httpx client)", "Switched to create_ssrf_safe_client"],
            ["H-2", "High", "SSRF via Slack webhook dispatch", "Switched to create_ssrf_safe_client"],
            ["H-3", "High", "SSRF via Jira client (3 callsites)", "Switched to create_ssrf_safe_client"],
            ["H-4", "High", "/register endpoint allowed viewer role", "Changed dependency to AdminUser"],
            ["H-5", "High", "/reset-password had no rate limiting", "Added 10/hour limit"],
            ["H-6", "High", "hmac.new does not exist in Python 3.8+", "Fixed to hmac.HMAC"],
            ["H-7", "High", "Deploy via SSH as root with IP in public repo", "User changed to deploy, IP moved to secret"],
            ["H-8", "High", "tenant_id nullable and never populated", "Backfilled and set NOT NULL"],
        ],
    )

    add_heading("Remaining Findings (Medium)", level=2)
    add_table(
        ["Finding", "Risk", "Recommended Action"],
        [
            ["/metrics endpoint publicly accessible", "Info disclosure", "Restrict to internal IPs via Caddy"],
            ["Invitation token in API response body", "Token leakage via logs", "Remove from response, deliver via email only"],
            ["CSP uses unsafe-inline for styles", "CSS injection", "Use nonce-based CSP"],
            ["HSTS missing preload directive", "First-visit downgrade", "Add preload, register at hstspreload.org"],
            ["Failed login attempts not audited", "Forensic gap", "Add audit log entry on failed auth"],
            ["SSO endpoints missing rate limit", "DoS on IdP discovery", "Added 20/minute limit"],
        ],
    )

    add_heading("Security Architecture Strengths", level=2)
    add_body(
        "JWT implementation: refresh token rotation with DB-backed revocation, MFA pending token "
        "separation, timing-safe dummy bcrypt on failed lookups, idle timeout enforcement. "
        "SSRF protection: double-layer with validate_public_url at input time and DNS rebinding "
        "protection at connection time via _SsrfSafeTransport. "
        "Credential storage: Fernet encryption at rest with key rotation capability. "
        "Rate limiting: proxy-aware IP extraction from X-Forwarded-For on all sensitive endpoints."
    )

    # ── 4. CODE QUALITY ──
    add_heading("4. Code Quality Assessment", level=1)
    add_body(
        "The codebase follows a clean layered architecture: API routes -> services -> models. "
        "No circular dependencies were found. FastAPI routers are thin, business logic resides "
        "in dedicated service modules, and SQLAlchemy models use modern 2.x mapped_column syntax."
    )
    add_table(
        ["Aspect", "Status", "Notes"],
        [
            ["Linting", "Clean", "ruff check + ruff format enforced in CI"],
            ["Type checking", "Clean", "mypy + TypeScript --noEmit pass"],
            ["Import ordering", "Automated", "ruff I (isort-compatible)"],
            ["Error handling", "Consistent", "HTTPException with status codes, no bare except"],
            ["Logging", "stdlib logging", "No print() statements, structured JSON in prod"],
            ["Evaluator", "Batched", "500 assets per batch, peak ~50 MB instead of full account in RAM"],
            ["Seed controls", "Atomic", "INSERT ON CONFLICT DO UPDATE, safe for concurrent startup"],
            ["Pagination", "Cursor-based", "Timeline, comments, evidence use CursorMeta with has_more"],
        ],
    )

    # ── 5. SCALABILITY ──
    add_heading("5. Scalability Assessment", level=1)
    add_table(
        ["Component", "Current Capacity", "Bottleneck", "Mitigation"],
        [
            ["Evaluator", "10,000+ assets/account", "RAM", "Batched to 500/iteration"],
            ["Evidence table", "Millions of rows", "Disk, query time", "90-day cleanup + cursor pagination"],
            ["Refresh tokens", "Unlimited", "Table bloat", "Hourly cleanup beat task"],
            ["Concurrent scans", "2 (Celery workers)", "Queue depth", "Increase concurrency or add workers"],
            ["Redis", "256 MB", "LRU eviction", "Separate instances recommended for scale"],
            ["Database connections", "30 max (pool)", "Connection exhaustion", "Sufficient for 50+ tenants"],
        ],
    )

    # ── 6. TEST COVERAGE ──
    add_heading("6. Test Coverage", level=1)
    add_table(
        ["Category", "Count", "Framework"],
        [
            ["Backend unit tests", "~600", "pytest (check functions, priority, remediation renderer)"],
            ["Backend integration tests", "~700", "pytest + httpx AsyncClient + ASGITransport"],
            ["Frontend unit tests", "89", "vitest + @testing-library/react"],
            ["E2E tests", "50", "Playwright"],
            ["Total", "1,297 backend + 89 frontend", ""],
        ],
    )

    # ── 7. DEPENDENCY HEALTH ──
    add_heading("7. Dependency Health", level=1)
    add_body(
        "All Dependabot security advisories were addressed on 12 April 2026. "
        "The alert count was reduced from 11 to 1 (a moderate-severity transitive dependency)."
    )
    add_table(
        ["Package", "Previous", "Updated To", "Severity"],
        [
            ["axios", "1.13.6", "1.15.0", "2 Critical"],
            ["next", "16.2.1", "16.2.3", "2 High"],
            ["lodash", "4.17.23", "4.18.1 (override)", "1 High + 1 Medium"],
            ["vite", "8.0.0", "8.0.8 (direct dev dep)", "2 High + 1 Medium"],
            ["picomatch", "2.3.1", "4.0.4 (override)", "2 Medium"],
        ],
    )

    # ── 8. PRODUCTION DEPLOYMENT ──
    add_heading("8. Production Deployment", level=1)
    add_table(
        ["Aspect", "Configuration"],
        [
            ["Host", "DigitalOcean VPS (single instance)"],
            ["Containers", "6 services via Docker Compose (backend, frontend, celery, beat, db, redis)"],
            ["Reverse proxy", "Caddy (host-level systemd service, automatic Let's Encrypt TLS)"],
            ["Deploy method", "GitHub Actions CD: SSH as 'deploy' user, git pull, docker compose up --build"],
            ["Migrations", "Alembic upgrade head runs in entrypoint.sh before uvicorn starts"],
            ["Monitoring", "Prometheus /metrics + /health endpoint (503 on DB/Redis degraded)"],
            ["Backups", "backup-db.sh script (manual scheduling required)"],
            ["Port binding", "All services on 127.0.0.1 only (no external exposure)"],
            ["Redis hardening", "Password required, dangerous commands renamed, maxmemory 256 MB"],
        ],
    )

    # ── 9. RECOMMENDATIONS ──
    add_heading("9. Recommendations", level=1)

    add_heading("Immediate (before next customer)", level=2)
    add_table(
        ["#", "Action", "Effort"],
        [
            ["1", "Configure SMTP (Resend) for password reset email delivery", "10 min (blocked on DNS)"],
            ["2", "Resolve remaining 1 moderate Dependabot alert", "15 min"],
            ["3", "Add automated DB backup cron (daily pg_dump to S3/object storage)", "1 hour"],
        ],
    )

    add_heading("Short-term (next 30 days)", level=2)
    add_table(
        ["#", "Action", "Effort"],
        [
            ["1", "Separate Redis instances for broker vs cache vs rate limiter", "2 hours"],
            ["2", "Add partial unique index on scans to prevent concurrent scan race", "1 hour"],
            ["3", "Extend remediation snippets from 40 to 100+ controls", "3 days (incremental)"],
            ["4", "Onboarding wizard: copy-paste Azure CLI for service principal creation", "1 day"],
        ],
    )

    add_heading("Medium-term (next 90 days)", level=2)
    add_table(
        ["#", "Action", "Effort"],
        [
            ["1", "AWS coverage expansion from 20 to 100+ checks", "2 weeks"],
            ["2", "Entra ID / Microsoft Graph collector (MFA controls CIS-AZ-01, 02)", "1 week"],
            ["3", "Multi-region deployment (EU/US data residency)", "2 weeks"],
            ["4", "Horizontal Celery worker scaling for concurrent scans", "1 week"],
        ],
    )

    # ── FOOTER ──
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("End of Report")
    run.font.color.rgb = GRAY
    run.font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Copyright 2026 Andrea Ceresoni. All rights reserved. Confidential.")
    run.font.color.rgb = GRAY
    run.font.size = Pt(8)

    doc.save(output_path)
    print(f"  DOCX generated: {output_path}")


if __name__ == "__main__":
    base = os.path.dirname(os.path.abspath(__file__))
    generate_product_pdf(os.path.join(base, "PostureOne_Product_Overview.pdf"))
    generate_technical_docx(os.path.join(base, "PostureOne_Technical_Assessment.docx"))
    print("Done.")
